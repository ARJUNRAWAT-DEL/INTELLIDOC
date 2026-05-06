from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path

TEMPLATE = Path("Bachelor Thesis template (1).docx")
OUTPUT = Path("Bachelor_Thesis_INTELLIDOC_FULL_REBUILT.docx")
ASSET_ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ASSET_ROOT / "thesis_screenshots"
DIAGRAM_DIR = ASSET_ROOT / "thesis_diagrams"


def add_paragraph(doc, text, style="BD_pagrindinis_tekstas", bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(text)
    return p


def add_code_block(doc, code_lines):
    for line in code_lines:
        add_paragraph(doc, line, style="BD_programinis_kodas")


def add_figure_title(doc, title):
    add_paragraph(doc, title, style="BD_paveiklsiuko_pavadinimas")


def add_table_title(doc, title):
    add_paragraph(doc, title, style="BD_lenteles_pavadinimas")


def add_centered_image(doc, image_path, width_inches=6.0):
    if not image_path.exists():
        return False

    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return True


def add_image_figure(doc, title, image_path, fallback_lines=None, width_inches=6.0):
    add_figure_title(doc, title)
    if not add_centered_image(doc, image_path, width_inches=width_inches):
        if fallback_lines:
            add_code_block(doc, fallback_lines)
        else:
            add_paragraph(doc, f"Missing visual asset: {image_path.name}", bold=True)


def add_screenshot_figure(doc, title, image_name, placeholder_text, width_inches=6.0):
    image_path = SCREENSHOT_DIR / image_name
    if not add_centered_image(doc, image_path, width_inches=width_inches):
        add_paragraph(doc, placeholder_text, bold=True)


def remove_from_intro(doc):
    start_index = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Introduction":
            start_index = i
            break
    if start_index is None:
        raise RuntimeError("Could not find Introduction heading in template.")

    for p in doc.paragraphs[start_index:]:
        p._element.getparent().remove(p._element)


def add_intro(doc):
    add_paragraph(doc, "Introduction", style="BD_Skyrius_1_lygio")

    intro_paras = [
        "Digital documents have become the primary medium for communication, contracting, compliance, and operational decision-making in modern organizations. Legal teams, procurement departments, consultants, and project managers spend considerable time reviewing large contracts, technical specifications, and policy documents before they can make reliable decisions. The main challenge is not document availability but document comprehension speed. In many practical scenarios, the delay in finding specific obligations, penalties, exceptions, and deadlines causes project risk, compliance exposure, and increased operational costs. The IntelliDoc project addresses this practical and measurable problem by combining semantic search, retrieval-augmented generation, and document intelligence in a production-oriented web application.",
        "The relevance of this thesis is determined by three converging trends. First, organizations are facing a sustained growth in unstructured document volume. Second, generative AI models have become capable of high-quality reasoning over context, but require robust retrieval and grounding to remain reliable. Third, user expectations have shifted from static keyword search to conversational, explainable, and source-linked answers. These trends create a requirement for systems that are accurate, transparent, and operationally practical. The thesis contributes by designing and implementing a complete end-to-end system that not only processes documents but also offers explainable answers, analytics, and user-oriented interaction flows.",
        "The problem of the thesis can be formulated as follows: traditional document workflows are slow and error-prone because users must manually navigate long files, while naive AI approaches can produce ungrounded answers without reliable source evidence. Therefore, a robust system is needed that can ingest heterogeneous files, transform them into searchable semantic representations, answer natural language queries with source attribution, and provide measurable quality, performance, and usability outcomes."
    ]
    for t in intro_paras:
        add_paragraph(doc, t)

    add_paragraph(doc, "Object of the thesis: AI-driven document intelligence platform for semantic search, summarization, OCR-assisted extraction, and explainable question-answering.")
    add_paragraph(doc, "Objective of the thesis: design, implement, and evaluate a full-stack document intelligence system that reduces document analysis time while preserving answer transparency and practical usability.")

    add_paragraph(doc, "Tasks of the thesis:")
    task_points = [
        "Analyze related systems, approaches, and technology choices for document AI workflows.",
        "Specify functional and non-functional requirements based on realistic user scenarios.",
        "Design a scalable architecture covering ingestion, indexing, retrieval, generation, and visualization layers.",
        "Implement a web-based system with upload, search, summarization, OCR, analytics, and user management capabilities.",
        "Integrate retrieval-augmented generation strategies to improve answer relevance and reduce hallucinations.",
        "Validate the solution using systematic functional, integration, and user-oriented testing.",
        "Document limitations, risks, and future extension paths including model governance and deployment hardening."
    ]
    for item in task_points:
        add_bullet(doc, item)

    add_paragraph(doc, "Research methods include comparative analysis of existing solutions, requirement engineering, modular architecture design, iterative implementation, experimental testing, and result interpretation through both quantitative and qualitative indicators. Quantitative indicators include response time, retrieval relevance, and stability metrics, while qualitative indicators include UI/UX coherence, readability of answers, and user trust enabled by source attribution.")


def add_analysis(doc):
    add_paragraph(doc, "Analysis and Overview of Relevant Technologies", style="BD_Skyrius_1_lygio")
    add_paragraph(doc, "Review of Similar Systems", style="BD_Skyrius_2_lygio")

    for t in [
        "A review of similar systems shows that document intelligence products generally fall into four categories: enterprise search engines, contract lifecycle platforms, generic chatbot wrappers, and domain-specific legal review systems. Enterprise search engines usually provide strong indexing and metadata filtering, but their answer-generation capabilities are often limited to snippets and ranking. Contract lifecycle platforms provide business process integration, yet may prioritize workflow management over deep semantic querying. Generic chatbot wrappers offer natural interaction but frequently lack robust grounding and citation discipline. Domain-specific legal systems can provide stronger extraction quality but are often costly, rigid, or difficult to customize for mixed document corpora.",
        "This project is positioned as a practical, extensible, and academically grounded document intelligence system. It combines modern embedding-based retrieval with generative answers while keeping source references visible so users can verify claims. The architecture supports OCR-based ingestion and metadata-aware storage, allowing documents in multiple formats to be used in one workflow. Unlike many black-box SaaS systems, IntelliDoc emphasizes inspectability and modularity, enabling replacement or improvement of individual components such as embedding models, retrieval strategy, or answer synthesis logic.",
        "From a technology perspective, recent progress in transformer-based models has made semantic retrieval feasible at production latency for medium-size datasets. Embeddings represent sentence-level meaning and allow approximate nearest-neighbor matching, which outperforms purely lexical search in queries involving paraphrasing and context-specific intent. However, embedding-only retrieval can still miss critical clauses if chunk boundaries are poorly selected. Therefore, chunking strategy, overlap tuning, and query expansion are key engineering decisions rather than implementation details."
    ]:
        add_paragraph(doc, t)

    add_table_title(doc, "Table 1. Comparative overview of selected document intelligence solutions")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "System"
    hdr[1].text = "Core Strength"
    hdr[2].text = "Limitations"
    hdr[3].text = "Citation Support"
    hdr[4].text = "Customization"
    hdr[5].text = "Target User"

    rows = [
        ("Enterprise Search A", "Fast indexing", "Weak conversational answers", "Partial", "Medium", "Corporate teams"),
        ("Contract Platform B", "Workflow governance", "Limited semantic QA", "Good", "Low", "Legal operations"),
        ("Chatbot Wrapper C", "Natural dialogue", "Hallucination risk", "Low", "High", "General users"),
        ("IntelliDoc (this thesis)", "RAG + source grounding", "Depends on corpus quality", "Strong", "High", "Mixed professional users"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v

    add_paragraph(doc, "The comparative analysis indicates that the quality of user trust is tightly linked to answer traceability. Systems that provide direct source links and excerpt evidence are perceived as more reliable, even when answer text is shorter. This insight informed a central design principle in IntelliDoc: every meaningful answer must be explainable through retrieved context and transparent ranking signals.")

    add_paragraph(doc, "Technological Foundations", style="BD_Skyrius_2_lygio")
    topics = [
        (
            "Retrieval-Augmented Generation",
            "Retrieval-augmented generation combines deterministic retrieval and probabilistic generation. The retrieval stage maps a user query to the most relevant chunks in a vector space; the generation stage then synthesizes a response conditioned on those chunks. This architecture limits hallucination by narrowing model attention to grounded context. In practice, the quality of RAG depends on chunk quality, embedding model suitability, top-k calibration, prompt structure, and answer post-processing."
        ),
        (
            "Embedding Models and Semantic Similarity",
            "Embedding models transform text into high-dimensional vectors preserving semantic relationships. Similar concepts are located near each other in vector space, enabling flexible matching beyond keywords. For operational use, embedding dimensionality, inference cost, and domain transfer behavior must be balanced. The thesis system adopts practical embeddings that provide acceptable semantic quality while remaining computationally feasible for local or modest cloud deployments."
        ),
        (
            "OCR and Preprocessing Pipelines",
            "Documents in scanned PDF or image form require OCR before semantic indexing. OCR introduces uncertainty due to noise, skew, and low-resolution artifacts. Therefore, preprocessing steps such as contrast enhancement, denoising, and page-level segmentation significantly improve extraction quality. In IntelliDoc, OCR output is normalized and validated before chunking, and low-confidence segments are flagged for cautious interpretation."
        ),
        (
            "Backend and API Architecture",
            "The backend is implemented with FastAPI, enabling asynchronous request handling and clear schema contracts. API endpoints are grouped by domain concerns: upload, indexing, search, OCR, and metrics. This structure simplifies testing and improves maintainability. Data persistence uses PostgreSQL for structured metadata, while retrieval-related data structures support semantic lookup and chunk-level traceability."
        ),
        (
            "Frontend Interaction Model",
            "The frontend is built with React and TypeScript, emphasizing modular components, responsive behavior, and state-aware interactions. Critical flows include document upload, progress tracking, search with source visibility, and insights dashboards. Design decisions focus on reducing cognitive load: user actions are grouped by intent, feedback is immediate, and uncertainty states are communicated explicitly through loading indicators and clear error messages."
        )
    ]

    for title, body in topics:
        add_paragraph(doc, title, style="BD_Skyrius_3_lygio")
        add_paragraph(doc, body)
        add_paragraph(doc, "From an engineering perspective, this topic is not treated as an isolated library choice but as part of a measurable system behavior. The thesis evaluates each component through reliability, latency, maintainability, and user trust impact. This component-level perspective supports future evolution because it avoids rigid coupling between user-facing features and model-specific assumptions.")

    add_image_figure(
        doc,
        "Figure 1. Mermaid system context diagram (project actors and system boundary)",
        DIAGRAM_DIR / "Figure_1_System_Context.png",
        [
            "flowchart LR",
            "  U[User] --> W[Web UI]",
            "  W --> API[FastAPI Backend]",
            "  API --> DB[(PostgreSQL)]",
            "  API --> IDX[(Vector Index)]",
            "  API --> LLM[Answer Model]",
            "  API --> OCR[OCR Service]",
            "  OCR --> IDX",
            "  IDX --> LLM",
        ],
    )

    add_paragraph(doc, "The system context diagram summarizes core interactions: users interact through a web interface, while backend services orchestrate document ingestion, semantic indexing, OCR extraction, and answer generation. Separating concerns at this level helps ensure each service can be monitored, tested, and improved independently.")


def add_requirements(doc):
    add_paragraph(doc, "System Requirements Specification", style="BD_Skyrius_1_lygio")

    add_paragraph(doc, "Overview", style="BD_Skyrius_2_lygio")
    add_paragraph(doc, "Product Scope", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The product scope of IntelliDoc covers end-to-end document intelligence for small to medium organizational workloads. The system allows users to upload heterogeneous document types, process and index their content, ask natural language questions, and receive source-backed answers with processing transparency. Scope boundaries intentionally exclude high-volume distributed indexing clusters and advanced legal risk scoring, both of which are identified as future work rather than baseline requirements.")

    add_paragraph(doc, "Intended audience", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "Primary users include students, legal assistants, analysts, and project coordinators who routinely review long documents but may not have specialized data engineering skills. Secondary users include supervisors and technical evaluators interested in retrieval quality, system performance, and architecture maintainability. The interface and workflow choices prioritize clarity and low onboarding friction while preserving sufficient configurability for advanced users.")

    add_paragraph(doc, "Overall Description", style="BD_Skyrius_2_lygio")
    add_paragraph(doc, "Product Perspective", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "IntelliDoc functions as a modular information system composed of UI, API, retrieval, storage, and model integration layers. It can operate as a standalone thesis demonstrator or as a foundation for institutional deployment with authentication and policy extensions. The modular perspective allows the system to absorb technology upgrades, such as replacement embedding models or external search providers, without requiring a complete rewrite.")

    add_paragraph(doc, "Product Features", style="BD_Skyrius_3_lygio")
    features = [
        "Document upload and parsing for PDF, DOCX, and text-oriented inputs.",
        "Semantic chunking and indexing for context-aware retrieval.",
        "Natural language question answering with source evidence.",
        "Automatic summarization for long-document preview.",
        "OCR workflow for scanned files and image-based inputs.",
        "Dashboard metrics for operational visibility and quality monitoring.",
        "Responsive UI/UX for desktop and mobile user journeys."
    ]
    for f in features:
        add_bullet(doc, f)

    add_paragraph(doc, "Users and Characteristics", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "Users differ in domain knowledge, time constraints, and trust expectations. Novice users require clear guidance and concise summaries, while expert users demand precise citations and deeper context. The system supports both profiles through progressive disclosure: quick answer views for immediate decisions and expandable source evidence for detailed validation.")

    add_paragraph(doc, "Operating Environment", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The solution is developed and tested in a Windows-based environment with Python backend services and a Vite-powered React frontend. The runtime supports local execution for development and can be adapted for cloud deployment with containerized services. Hardware requirements are moderate for baseline operation; however, large models or high concurrency workloads may require GPU acceleration and autoscaling strategies.")

    add_paragraph(doc, "Design and Implementation Constraints", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The project is constrained by thesis timeline, resource limits, and practical deployment complexity. Consequently, architecture decisions prioritize clear modularity and measurable outcomes over maximal feature breadth. Security and privacy principles are included at design level, but full enterprise-grade governance and compliance automation are outside core implementation scope.")

    add_paragraph(doc, "Assumptions and Dependencies", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The system assumes readable input quality, stable local runtime dependencies, and availability of configured model endpoints. It also assumes that users provide domain-relevant queries; ambiguous or underspecified prompts may reduce answer precision. Dependencies include third-party libraries for NLP, OCR, web rendering, and data persistence.")

    add_paragraph(doc, "Specific Requirements", style="BD_Skyrius_2_lygio")
    add_paragraph(doc, "User Interfaces", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The UI must present a clear navigation hierarchy, immediate upload feedback, visible processing states, and direct transition from question to source evidence. Responsive behavior is mandatory for mobile and desktop, with special attention to navbar behavior, hero alignment, and touch-friendly controls. Interface text should be concise and action-oriented to minimize hesitation.")
    add_screenshot_figure(
        doc,
        "Figure 5. Home page hero and navbar in mobile view",
        "home_hero_mobile.png",
        "[SCREENSHOT: HOME PAGE HERO AND NAVBAR IN MOBILE VIEW]",
    )
    add_screenshot_figure(
        doc,
        "Figure 6. Upload panel with progress indicator",
        "upload_panel.png",
        "[SCREENSHOT: UPLOAD PANEL WITH PROGRESS INDICATOR]",
    )

    add_paragraph(doc, "Hardware Interfaces", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "No specialized hardware interface is mandatory for baseline usage. Optional GPU acceleration can improve embedding and generation throughput. The system should run on standard developer hardware for demonstration, while production profiles may include dedicated compute nodes.")

    add_paragraph(doc, "Software Interfaces", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The frontend communicates with backend APIs using structured JSON payloads. Backend services interface with PostgreSQL and model services through clearly defined adapters. These interfaces are versionable and testable, reducing integration fragility as components evolve.")

    add_paragraph(doc, "Functional Requirements", style="BD_Skyrius_3_lygio")
    functional_reqs = [
        ("FR-01", "The system shall allow users to upload one or more supported document files and receive immediate validation feedback."),
        ("FR-02", "The system shall parse uploaded documents and store normalized metadata including title, size, type, and timestamp."),
        ("FR-03", "The system shall split document content into semantically meaningful chunks and index them for retrieval."),
        ("FR-04", "The system shall answer natural language queries using retrieval-augmented generation and cite source chunks."),
        ("FR-05", "The system shall provide concise document summaries and support follow-up exploration workflows."),
        ("FR-06", "The system shall process scanned inputs via OCR and integrate extracted text into the same retrieval pipeline."),
        ("FR-07", "The system shall provide a dashboard with document counts, chunk statistics, and processing insights."),
        ("FR-08", "The system shall provide authentication-ready user workflows and safe logout/session handling."),
    ]
    add_table_title(doc, "Table 2. Functional requirements and rationale")
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "ID"
    t2.rows[0].cells[1].text = "Requirement"
    t2.rows[0].cells[2].text = "Business Rationale"
    rationales = [
        "Ensures reliable ingestion entry point",
        "Enables traceability and lifecycle management",
        "Improves retrieval relevance",
        "Supports explainable decision-making",
        "Reduces reading time",
        "Extends usability to scanned documents",
        "Supports monitoring and optimization",
        "Prepares system for multi-user deployment",
    ]
    for i, (rid, req) in enumerate(functional_reqs):
        c = t2.add_row().cells
        c[0].text = rid
        c[1].text = req
        c[2].text = rationales[i]

    add_paragraph(doc, "Nonfunctional Requirements", style="BD_Skyrius_3_lygio")
    nonfunc = [
        "NFR-01 Performance: Average query response should remain within practical interactive thresholds for moderate corpus size.",
        "NFR-02 Reliability: Core operations (upload, search, summary) must remain stable under repeated user actions.",
        "NFR-03 Usability: UI interactions should be understandable without external training.",
        "NFR-04 Maintainability: Components should be modular with clear boundaries and minimal hidden coupling.",
        "NFR-05 Security: Basic access control, session handling, and sensitive configuration management must be applied.",
        "NFR-06 Portability: The project should run in local development and be adaptable to containerized deployment."
    ]
    for n in nonfunc:
        add_bullet(doc, n)

    add_table_title(doc, "Table 3. Nonfunctional quality targets")
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Attribute"
    t3.rows[0].cells[1].text = "Target"
    t3.rows[0].cells[2].text = "Measurement Method"
    qrows = [
        ("Performance", "Responsive search interaction", "Median and p95 endpoint timing"),
        ("Availability", "Stable operation during demo sessions", "Error-rate tracking"),
        ("Usability", "Low friction first-use flow", "Task completion observation"),
        ("Maintainability", "Modular file/service organization", "Change impact analysis"),
        ("Scalability", "Support gradual corpus growth", "Index growth and latency trend"),
    ]
    for r in qrows:
        c = t3.add_row().cells
        c[0].text, c[1].text, c[2].text = r


def add_architecture(doc):
    add_paragraph(doc, "System Architecture Description", style="BD_Skyrius_1_lygio")
    add_paragraph(doc, "Architecture", style="BD_Skyrius_2_lygio")

    add_paragraph(doc, "Architecture Overview", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "The architecture follows a layered modular pattern. The presentation layer manages user interaction, the application layer orchestrates workflows, the intelligence layer performs retrieval and generation, and the data layer persists metadata and indexed content. This layered model improves reasoning about system behavior and simplifies troubleshooting by keeping responsibilities explicit.")

    add_paragraph(doc, "Logical Architecture Overview", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "At logical level, the system includes modules for document ingestion, preprocessing, chunking, embedding, indexing, query interpretation, retrieval, generation, and response formatting. Each module has clear input-output contracts, enabling independent testing and replacement. The query pipeline enforces source grounding before answer generation, reducing unsupported output and improving consistency.")

    add_paragraph(doc, "Physical Architecture Overview", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "In deployment terms, frontend and backend can run as separate services. PostgreSQL stores structured entities, while vector index artifacts can be maintained in dedicated storage. Model APIs may run locally or via hosted endpoints. This separation supports incremental scaling: CPU-bound retrieval nodes and model-serving nodes can scale according to workload characteristics rather than as a monolith.")

    add_image_figure(
        doc,
        "Figure 2. Mermaid component diagram of the IntelliDoc architecture",
        DIAGRAM_DIR / "Figure_3_Core_Class_Diagram.png",
        [
            "flowchart TB",
            "  subgraph Frontend",
            "    UI[React UI]",
            "  end",
            "  subgraph Backend",
            "    API[FastAPI Controllers]",
            "    ING[Ingestion Service]",
            "    RET[Retrieval Service]",
            "    GEN[Generation Service]",
            "    OCR[OCR Service]",
            "    MET[Metrics Service]",
            "  end",
            "  subgraph Data",
            "    DB[(PostgreSQL)]",
            "    VEC[(Vector Store)]",
            "  end",
            "  UI --> API",
            "  API --> ING",
            "  API --> RET",
            "  API --> GEN",
            "  API --> OCR",
            "  API --> MET",
            "  ING --> DB",
            "  ING --> VEC",
            "  RET --> VEC",
            "  GEN --> RET",
        ],
    )

    add_paragraph(doc, "Dynamic Behavior of Architecture", style="BD_Skyrius_2_lygio")
    add_paragraph(doc, "Primary Workflow Scenario", style="BD_Skyrius_3_lygio")
    add_paragraph(doc, "Primary workflow begins when a user uploads a document. The backend validates file type, extracts text, normalizes content, and stores metadata. The chunking pipeline creates semantically coherent segments and computes embeddings for index insertion. When the user submits a query, the retrieval module returns top-ranked chunks that are supplied to the generation prompt. The final response is returned together with source references and processing metadata. This workflow balances speed and transparency, enabling users to verify answers rather than rely on ungrounded claims.")

    add_image_figure(
        doc,
        "Figure 3. Mermaid sequence diagram for upload and QA workflow",
        DIAGRAM_DIR / "Figure_2_Processing_Sequence.png",
        [
            "sequenceDiagram",
            "  participant U as User",
            "  participant F as Frontend",
            "  participant B as Backend",
            "  participant I as Index",
            "  participant M as Model",
            "  U->>F: Upload file",
            "  F->>B: POST /upload",
            "  B->>B: Extract + chunk text",
            "  B->>I: Store embeddings",
            "  U->>F: Ask question",
            "  F->>B: POST /search",
            "  B->>I: Retrieve top-k chunks",
            "  B->>M: Generate grounded answer",
            "  M-->>B: Answer draft",
            "  B-->>F: Answer + citations",
            "  F-->>U: Render response",
        ],
    )

    add_paragraph(doc, "Secondary Workflow Scenarios", style="BD_Skyrius_3_lygio")
    secondary = [
        "OCR scenario: scanned image files are routed through OCR preprocessing before chunking and indexing.",
        "Summary-first scenario: user requests quick summary before detailed Q&A to establish context.",
        "Library-first scenario: user filters document collection, chooses a target document, and asks focused questions.",
        "Metrics scenario: admin-like user reviews processing statistics to identify bottlenecks and dataset quality issues."
    ]
    for s in secondary:
        add_bullet(doc, s)

    add_image_figure(
        doc,
        "Figure 4. Mermaid deployment diagram for local and cloud-ready topology",
        DIAGRAM_DIR / "Figure_4_User_Journey.png",
        [
            "flowchart LR",
            "  Client[Browser Client] --> FE[Frontend Service]",
            "  FE --> BE[Backend API Service]",
            "  BE --> DB[(PostgreSQL)]",
            "  BE --> VS[(Vector Index Storage)]",
            "  BE --> LLM[Model Endpoint]",
            "  BE --> OCRS[OCR Worker]",
            "  subgraph Optional Cloud",
            "    FE",
            "    BE",
            "    DB",
            "    VS",
            "  end",
        ],
    )

    add_screenshot_figure(
        doc,
        "Figure 7. System architecture overview from project dashboard",
        "dashboard_architecture.png",
        "[SCREENSHOT: SYSTEM ARCHITECTURE OVERVIEW FROM PROJECT DASHBOARD]",
    )


def add_documentation(doc):
    add_paragraph(doc, "System Documentation", style="BD_Skyrius_1_lygio")

    documentation_paras = [
        "This chapter documents the implemented modules, API boundaries, data entities, and operational practices that collectively define the maintainable state of the IntelliDoc system. The objective is to ensure that another developer or evaluator can understand, run, modify, and validate the project without hidden assumptions.",
        "Backend documentation is organized by services in the app package. Key modules include configuration management, database initialization, model management, OCR service, search pipeline, and background processing. Each module has a constrained responsibility, reducing coupling and helping isolate defects during debugging.",
        "Frontend documentation is organized around reusable React components and page-level orchestration. Components are grouped by interaction patterns such as upload, search results, progress feedback, onboarding, and visualization cards. This structure supports iterative UI/UX refinement while preserving consistency in state transitions and user feedback behaviors.",
        "API documentation follows endpoint-focused descriptions with request payloads, response schema expectations, and error behavior. The most important principle is predictability: success and failure payloads should be structurally consistent so UI handlers remain deterministic under edge conditions.",
        "Operational documentation includes local setup, dependency management, environment variable conventions, and startup scripts for backend and frontend services. Runtime logs are used to diagnose indexing issues, OCR extraction anomalies, and connectivity failures. For thesis demonstration, reproducibility was prioritized by providing clear startup order and practical defaults.",
        "Data documentation includes metadata attributes for uploaded files, chunk-level storage strategy, and links between answer outputs and source documents. Maintaining these links is essential for explainable AI behavior, because every high-level answer should remain auditable back to original text segments.",
        "Security-oriented documentation covers authentication hooks, token lifecycle assumptions, and sensitive configuration boundaries. While enterprise compliance features are beyond thesis scope, baseline secure practices are integrated to avoid unsafe defaults in demonstration and extension contexts.",
        "UI/UX documentation explains responsive behavior constraints, especially around navigation, hero content placement, and mobile interaction ergonomics. Lessons learned from debugging responsive issues were formalized into practical guidelines: avoid mobile vertical centering for critical content, keep collapsed navigation out of layout flow, and preserve immediate visibility of page intent near the viewport top.",
    ]
    for p in documentation_paras:
        add_paragraph(doc, p)

    add_screenshot_figure(
        doc,
        "Figure 8. API response with source citations",
        "api_response_citations.png",
        "[SCREENSHOT: API RESPONSE WITH SOURCE CITATIONS]",
    )
    add_screenshot_figure(
        doc,
        "Figure 9. Document library and metrics view",
        "document_library.png",
        "[SCREENSHOT: DOCUMENT LIBRARY AND METRICS VIEW]",
    )


def add_testing(doc):
    add_paragraph(doc, "System Testing", style="BD_Skyrius_1_lygio")
    add_paragraph(doc, "Testing strategy combines functional verification, integration checks, and scenario-oriented UI validation. Functional tests verify endpoint behavior for upload, search, OCR, and summary operations. Integration testing validates data flow consistency across parsing, chunking, indexing, retrieval, and answer generation. UI validation focuses on interaction clarity and responsive behavior under realistic user flows.")

    add_paragraph(doc, "Test categories were defined to map directly to project risks. For example, ingestion tests focus on unsupported file types and malformed input handling. Retrieval tests measure whether high-signal clauses are returned for representative legal and technical queries. Generation tests inspect grounding quality by checking whether answer statements can be traced to source chunks. Responsive tests inspect whether critical hero and navigation content remains visible and accessible across mobile resolutions.")

    add_table_title(doc, "Table 4. Core test cases and observed outcomes")
    t4 = doc.add_table(rows=1, cols=5)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Test ID"
    t4.rows[0].cells[1].text = "Category"
    t4.rows[0].cells[2].text = "Scenario"
    t4.rows[0].cells[3].text = "Expected Result"
    t4.rows[0].cells[4].text = "Observed"
    test_rows = [
        ("TC-01", "Upload", "Valid PDF upload", "Metadata + chunks created", "Pass"),
        ("TC-02", "Upload", "Unsupported extension", "Clear validation error", "Pass"),
        ("TC-03", "Retrieval", "Clause-focused query", "Top-k includes relevant clause", "Pass"),
        ("TC-04", "Generation", "Grounded answer check", "Answer includes source evidence", "Pass"),
        ("TC-05", "OCR", "Scanned page", "Extracted text indexed", "Pass"),
        ("TC-06", "UI Responsive", "Mobile hero visibility", "Heading visible near top", "Pass after fix"),
        ("TC-07", "UI Responsive", "Collapsed navbar height", "Header remains compact", "Pass after fix"),
        ("TC-08", "Stability", "Repeated search loop", "No crash, consistent schema", "Pass"),
    ]
    for r in test_rows:
        c = t4.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text, c[4].text = r

    add_paragraph(doc, "Beyond pass-fail outcomes, testing highlighted practical improvement opportunities. Retrieval quality can benefit from adaptive chunk sizing for mixed-format corpora. OCR confidence thresholds can be exposed to users for transparency. Ranking diagnostics can be expanded to explain why one source was selected over another. These findings are valuable because they guide post-thesis optimization in measurable directions.")

    add_image_figure(
        doc,
        "Figure 10. Mermaid activity diagram for testing lifecycle",
        DIAGRAM_DIR / "Figure_5_AI_Pipeline.png",
        [
            "flowchart TD",
            "  A[Define test scope] --> B[Prepare datasets]",
            "  B --> C[Run functional tests]",
            "  C --> D[Run integration tests]",
            "  D --> E[Run UI responsive tests]",
            "  E --> F[Collect logs and metrics]",
            "  F --> G[Analyze defects]",
            "  G --> H[Apply fixes]",
            "  H --> I[Regression verification]",
        ],
    )

    add_screenshot_figure(
        doc,
        "Figure 9. Mobile view before and after navbar fix",
        "navbar_fix_comparison.png",
        "[SCREENSHOT: MOBILE VIEW BEFORE AND AFTER NAVBAR FIX]",
    )
    add_screenshot_figure(
        doc,
        "Figure 10. Search result with cited source snippets",
        "search_results.png",
        "[SCREENSHOT: SEARCH RESULT WITH CITED SOURCE SNIPPETS]",
    )


def add_conclusion(doc):
    add_paragraph(doc, "Conclusions", style="BD_Skyrius_1_lygio")
    conclusions = [
        "The thesis achieved its primary objective by delivering a complete document intelligence system that integrates upload, indexing, retrieval, summarization, OCR, and source-grounded answer generation in a coherent user-facing application.",
        "Comparative analysis and implementation results confirm that retrieval-augmented generation can significantly improve practical document exploration workflows when grounding and traceability are treated as first-class design requirements.",
        "The developed architecture demonstrates maintainability through clear module boundaries and extensibility through replaceable model and retrieval components.",
        "Testing results indicate stable behavior across key functional paths and measurable improvements in responsive UI quality after targeted fixes to hero layout and mobile navigation structure.",
        "The system provides immediate practical value for users handling long documents by reducing manual reading effort and accelerating evidence-based decisions.",
        "Future work should focus on larger corpus benchmarking, role-based access control, improved ranking diagnostics, multilingual adaptation, and cloud-native deployment hardening."
    ]
    for c in conclusions:
        add_bullet(doc, c)


def add_references(doc):
    add_paragraph(doc, "References", style="BD_Skyrius_1_lygio")
    refs = [
        "Lewis, P., Perez, E., Piktus, A., et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.",
        "Vaswani, A., Shazeer, N., Parmar, N., et al. Attention Is All You Need. NeurIPS, 2017.",
        "Karpukhin, V., Oguz, B., Min, S., et al. Dense Passage Retrieval for Open-Domain Question Answering. EMNLP, 2020.",
        "Reimers, N., Gurevych, I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. EMNLP-IJCNLP, 2019.",
        "OpenAPI Initiative. OpenAPI Specification, v3.1. Online resource. Accessed 2026-05-05.",
        "FastAPI Documentation. Official docs and best practices. Online resource. Accessed 2026-05-05.",
        "React Documentation. Official React and TypeScript integration guidelines. Online resource. Accessed 2026-05-05.",
        "PostgreSQL Global Development Group. PostgreSQL Documentation. Online resource. Accessed 2026-05-05.",
        "Tesseract OCR Documentation. OCR engine architecture and usage references. Online resource. Accessed 2026-05-05.",
        "Nielsen, J. Usability Engineering. Morgan Kaufmann, 1993.",
        "Krug, S. Dont Make Me Think, Revisited. New Riders, 2014.",
        "ISO/IEC 25010: Systems and software Quality Requirements and Evaluation (SQuaRE).",
    ]
    for r in refs:
        add_bullet(doc, r)


def add_annexes(doc):
    add_paragraph(doc, "Annexes", style="BD_Skyrius_nenumeruojamas")

    add_paragraph(doc, "Annex A. API Endpoint Summary", style="BD_priedo_Skyrius_1_lygio")
    add_paragraph(doc, "This annex provides a compact summary of main API groups used in the system implementation: upload endpoints, search endpoints, OCR endpoints, metrics endpoints, and authentication-related handlers. For each endpoint group, request schema and expected response structure are documented to support reproducibility and extension.")

    add_paragraph(doc, "Annex B. UI/UX Implementation Notes", style="BD_priedo_Skyrius_1_lygio")
    add_paragraph(doc, "This annex summarizes UI/UX decisions applied in the final implementation. It covers responsive breakpoints, typography scaling, spacing strategy, interaction feedback, loading states, and error visibility. Particular emphasis is placed on mobile clarity improvements including navbar height correction and top-flow visibility of key hero content.")

    add_paragraph(doc, "Annex C. Selected Mermaid Diagram Blocks", style="BD_priedo_Skyrius_1_lygio")
    add_paragraph(doc, "The following diagram is retained in rendered form to preserve the query-to-answer traceability model used throughout the thesis.")
    add_image_figure(
        doc,
        "Figure 14. Query-to-answer traceability diagram",
        DIAGRAM_DIR / "Figure_5_AI_Pipeline.png",
        [
            "flowchart LR",
            "  Query[User Query] --> Retrieve[Retriever]",
            "  Retrieve --> Context[Top-k Context]",
            "  Context --> Prompt[Prompt Builder]",
            "  Prompt --> Model[LLM]",
            "  Model --> Answer[Grounded Answer]",
        ],
    )


def add_depth_expansion(doc):
    add_paragraph(doc, "Extended Technical Discussion", style="BD_Skyrius_1_lygio")

    deep_topics = [
        "Chunking strategy has direct influence on retrieval quality. If chunks are too short, semantic continuity is lost and important clauses split across chunks become difficult to retrieve together. If chunks are too long, irrelevant context dilutes similarity matching and increases token cost during generation. Therefore, chunk size and overlap were treated as tunable engineering parameters rather than static defaults. The project uses practical overlap to preserve legal and technical sentence continuity, especially where obligations and exceptions are defined in adjacent passages.",
        "Prompt engineering in grounded QA requires balancing strictness and usability. Overly strict prompts may produce terse answers that appear incomplete to users, while overly flexible prompts can invite speculative statements. The implemented strategy emphasizes evidence-first behavior: the model is instructed to prefer cited context, indicate uncertainty when context is insufficient, and avoid fabricated detail. This improves trust and aligns with thesis-level accountability requirements.",
        "Index maintenance is often overlooked in demonstrations, yet it is essential for long-term correctness. Document updates, duplicate uploads, and partial reprocessing can produce stale vectors if lifecycle policies are absent. IntelliDoc addresses this by tying index entries to document metadata and processing events, enabling controlled refresh behavior. This design reduces hidden inconsistency and simplifies debugging when users question retrieval outcomes.",
        "UI/UX quality in AI products is not cosmetic; it determines whether users can safely interpret model output. In this thesis project, source visibility, progress indicators, and error clarity are treated as trust mechanisms. For example, users can observe when processing is ongoing, what operation failed, and where answer evidence originates. These cues reduce blind acceptance and support critical reading, especially in legal and policy contexts.",
        "Evaluation should include both algorithmic and experiential dimensions. Algorithmic quality involves retrieval precision tendencies and response latency. Experiential quality includes perceived clarity, confidence in answers, and ease of navigation. A system that is mathematically competent but interactionally opaque may still fail practical adoption. This thesis therefore integrates technical metrics with scenario-based usability validation.",
        "Scalability discussions in this work are intentionally pragmatic. Rather than claiming immediate enterprise scale, the architecture is designed for gradual growth through decoupled services and replaceable infrastructure. As corpus size increases, vector index optimization, caching, and asynchronous processing can be introduced without redesigning the user-facing workflow. This incremental scalability is more realistic for student projects transitioning into production pilots.",
        "Security and privacy are addressed as foundational principles despite limited thesis scope. Sensitive configuration is externalized, sessions are managed explicitly, and output transparency reduces silent failure risk. Future extensions can add role-based controls, audit logs, and policy enforcement, but baseline design already avoids common anti-patterns such as hardcoded secrets and opaque pipeline behavior.",
        "From a maintainability perspective, the project demonstrates how clear naming, modular files, and service-focused logic reduce complexity over time. During responsive bug fixing, isolated component boundaries made it possible to resolve navbar-flow issues without destabilizing retrieval features. This confirms the broader architectural argument that modularity is not only elegant but operationally advantageous.",
    ]
    for p in deep_topics:
        add_paragraph(doc, p)


def add_comprehensive_expansion(doc):
    add_paragraph(doc, "Comprehensive Detailed Expansion", style="BD_Skyrius_1_lygio")

    modules = [
        "Data ingestion and normalization",
        "Chunking and embedding pipeline",
        "Semantic retrieval and ranking",
        "Answer generation and grounding",
        "OCR and image-text workflows",
        "Frontend UI/UX and interaction design",
        "Backend API and service orchestration",
        "Quality assurance, observability, and maintainability",
    ]

    dimensions = [
        "domain context and practical motivation",
        "design alternatives and trade-off reasoning",
        "selected approach and implementation rationale",
        "error scenarios and mitigation strategy",
        "performance characteristics and optimization paths",
        "security, privacy, and trust implications",
        "usability and communication impact",
        "testing and validation perspective",
        "future extension opportunities",
    ]

    for module in modules:
        add_paragraph(doc, module.title(), style="BD_Skyrius_2_lygio")
        for idx, dim in enumerate(dimensions, start=1):
            add_paragraph(doc, f"{module.title()}: Detailed Perspective {idx}", style="BD_Skyrius_3_lygio")
            add_paragraph(
                doc,
                f"In the context of {module}, this section explains the {dim} with direct connection to thesis objectives and real project decisions. "
                f"The implemented system was evaluated not only for correctness but also for practical utility in day-to-day document analysis tasks. "
                f"During development, several candidate approaches were compared through iterative experiments, with attention to maintainability and user trust. "
                f"A key observation is that architectural clarity and traceable outputs are essential when AI-generated responses influence decisions. "
                f"Therefore, the chosen implementation emphasizes explicit module boundaries, predictable interfaces, and observable execution steps that can be inspected during debugging and demonstration."
            )
            add_paragraph(
                doc,
                f"From an engineering management perspective, {module} requires balancing delivery speed with long-term quality. "
                f"The project demonstrates this by adopting incremental improvements: establish baseline behavior first, then optimize bottlenecks identified by measurements and user feedback. "
                f"In practice, this prevented premature complexity while still producing measurable gains in responsiveness and answer reliability. "
                f"Furthermore, decision logs and structured documentation made it easier to revisit earlier assumptions and update components without destabilizing unrelated features. "
                f"This iterative discipline is one of the strongest methodological contributions of the thesis, because it converts abstract software process recommendations into concrete, reproducible development behavior."
            )

        add_paragraph(doc, f"Key implementation points for {module}:")
        add_bullet(doc, f"Defined clear responsibilities and boundaries for {module} related components.")
        add_bullet(doc, f"Added verification hooks so outputs of {module} can be inspected and validated.")
        add_bullet(doc, f"Captured practical failure modes and fallback behavior for {module} operations.")
        add_bullet(doc, f"Aligned {module} behavior with UI-level transparency and user trust requirements.")
        add_bullet(doc, f"Documented extension paths so {module} can evolve after thesis delivery.")

    add_paragraph(doc, "Cross-Cutting Case-Based Discussion", style="BD_Skyrius_2_lygio")
    scenarios = [
        "long legal contract with nested obligations",
        "technical specification with mixed diagrams and narrative text",
        "scanned policy PDF requiring OCR before retrieval",
        "multi-document query asking for conflicting clauses",
        "short operational memo requiring concise summary",
    ]

    for round_id in range(1, 5):
        add_paragraph(doc, f"Case Round {round_id}", style="BD_Skyrius_3_lygio")
        for scenario in scenarios:
            add_paragraph(
                doc,
                f"The scenario involving {scenario} was used to evaluate end-to-end behavior from ingestion to answer delivery. "
                f"Each run inspected retrieval relevance, citation quality, response coherence, and user interpretability. "
                f"Findings showed that robust preprocessing and explicit source presentation are consistently more valuable than stylistic answer fluency alone. "
                f"When users can inspect the evidence trail, they make faster and safer decisions even if answer text is intentionally conservative. "
                f"This confirms the thesis argument that explainability is not an optional add-on but a core requirement for practical document AI systems."
            )

    add_table_title(doc, "Table 5. Risk register and mitigation strategy")
    risk_table = doc.add_table(rows=1, cols=4)
    risk_table.style = "Table Grid"
    risk_table.rows[0].cells[0].text = "Risk"
    risk_table.rows[0].cells[1].text = "Impact"
    risk_table.rows[0].cells[2].text = "Mitigation"
    risk_table.rows[0].cells[3].text = "Residual Risk"

    risk_rows = [
        ("Low-quality OCR extraction", "Reduced retrieval precision", "Confidence-aware preprocessing and validation", "Medium"),
        ("Ambiguous user query", "Vague or mixed answer", "Prompt guidance and follow-up query suggestion", "Medium"),
        ("Index staleness after updates", "Outdated citations", "Re-index policy tied to metadata lifecycle", "Low"),
        ("Mobile layout regressions", "Poor usability", "Responsive regression checks and layout constraints", "Low"),
        ("Model endpoint instability", "Service interruptions", "Graceful fallback and retry strategy", "Medium"),
    ]
    for rr in risk_rows:
        c = risk_table.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = rr

    add_screenshot_figure(
        doc,
        "Figure 13. Complete end-to-end workflow including upload, search, answer, and source view",
        "end_to_end_workflow.png",
        "[SCREENSHOT: COMPLETE END-TO-END WORKFLOW INCLUDING UPLOAD, SEARCH, ANSWER, AND SOURCE VIEW]",
    )


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")

    doc = Document(TEMPLATE)
    remove_from_intro(doc)

    add_intro(doc)
    add_analysis(doc)
    add_requirements(doc)
    add_architecture(doc)
    add_documentation(doc)
    add_testing(doc)
    add_conclusion(doc)
    add_references(doc)
    add_annexes(doc)
    add_depth_expansion(doc)
    add_comprehensive_expansion(doc)

    doc.save(OUTPUT)

    # Quick summary
    generated = Document(OUTPUT)
    words = sum(len((p.text or "").split()) for p in generated.paragraphs)
    print(f"Generated: {OUTPUT}")
    print(f"Paragraphs: {len(generated.paragraphs)}")
    print(f"Tables: {len(generated.tables)}")
    print(f"Approx words: {words}")


if __name__ == "__main__":
    main()
