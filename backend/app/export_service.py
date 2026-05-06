from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any, Iterable, List, Optional

from .logger import logger

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
except Exception:  # pragma: no cover - handled at runtime
    DocxDocument = None
    Pt = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
except Exception:  # pragma: no cover - handled at runtime
    colors = None
    letter = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    inch = None
    SimpleDocTemplate = None
    Paragraph = None
    Spacer = None
    Table = None
    TableStyle = None


@dataclass
class ExportPayload:
    question: str
    answer: str
    document_name: str
    answer_length: str
    answer_mode: str
    page_range_label: str
    sources: List[Any]
    citations: List[Any]
    generated_at: datetime


def _safe_text(value: Any) -> str:
    return "" if value is None else str(value)


def _title_case(value: str, fallback: str) -> str:
    cleaned = (_safe_text(value) or "").strip()
    if not cleaned:
        return fallback
    mapping = {
        "short": "Short",
        "balanced": "Balanced",
        "detailed": "Detailed",
        "summary": "Summary",
        "qa": "Question Answering",
        "keypoints": "Key Points",
        "pageexplanation": "Page Explanation",
        "actionitems": "Action Items",
    }
    return mapping.get(cleaned.lower(), cleaned.replace("_", " ").title())


def _format_page_range(page_range: Optional[Any]) -> str:
    if not page_range:
        return "Entire document"
    start = getattr(page_range, "start", None) if not isinstance(page_range, dict) else page_range.get("start")
    end = getattr(page_range, "end", None) if not isinstance(page_range, dict) else page_range.get("end")
    if start and end:
        if start == end:
            return f"Page {start}"
        return f"Pages {start}-{end}"
    if start:
        return f"Page {start}"
    if end:
        return f"Page {end}"
    return "Entire document"


def build_export_payload(
    *,
    question: str,
    answer: str,
    sources: Optional[List[Any]] = None,
    citations: Optional[List[Any]] = None,
    document_name: Optional[str] = None,
    answer_length: str = "balanced",
    answer_mode: str = "summary",
    page_range: Optional[Any] = None,
    generated_at: Optional[datetime] = None,
) -> ExportPayload:
    doc_name = (document_name or "").strip()
    if not doc_name and sources:
        first_source = sources[0]
        doc_name = _safe_text(getattr(first_source, "doc_title", None) if not isinstance(first_source, dict) else first_source.get("doc_title"))
    if not doc_name:
        doc_name = "IntelliDoc Answer"

    return ExportPayload(
        question=_safe_text(question).strip(),
        answer=_safe_text(answer).strip(),
        document_name=doc_name,
        answer_length=_title_case(answer_length, "Balanced"),
        answer_mode=_title_case(answer_mode, "Summary"),
        page_range_label=_format_page_range(page_range),
        sources=list(sources or []),
        citations=list(citations or []),
        generated_at=generated_at or datetime.utcnow(),
    )


def _citation_label(citation: Any) -> str:
    page_number = getattr(citation, "page_number", None) if not isinstance(citation, dict) else citation.get("page_number")
    paragraph_number = getattr(citation, "paragraph_number", None) if not isinstance(citation, dict) else citation.get("paragraph_number")
    quote = getattr(citation, "quote", None) if not isinstance(citation, dict) else citation.get("quote")
    doc_title = getattr(citation, "doc_title", None) if not isinstance(citation, dict) else citation.get("doc_title")
    parts = []
    if doc_title:
        parts.append(_safe_text(doc_title))
    if page_number:
        parts.append(f"Page {page_number}")
    if paragraph_number:
        parts.append(f"Paragraph {paragraph_number}")
    if not parts:
        parts.append("Citation")
    label = ", ".join(parts)
    if quote:
        label = f"{label}: {_safe_text(quote)}"
    return label


def _source_label(source: Any) -> str:
    doc_title = getattr(source, "doc_title", None) if not isinstance(source, dict) else source.get("doc_title")
    doc_id = getattr(source, "doc_id", None) if not isinstance(source, dict) else source.get("doc_id")
    if doc_title and doc_id is not None:
        return f"{doc_title} (Document ID: {doc_id})"
    if doc_title:
        return _safe_text(doc_title)
    if doc_id is not None:
        return f"Document ID: {doc_id}"
    return "Unknown source"


def generate_pdf_bytes(payload: ExportPayload) -> bytes:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab is not installed")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=48, leftMargin=48, topMargin=48, bottomMargin=48)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="AnswerHeading", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=12, textColor=colors.HexColor("#0f172a")))
    styles.add(ParagraphStyle(name="MetaLabel", parent=styles["Normal"], fontSize=10, leading=13, textColor=colors.HexColor("#475569")))
    styles.add(ParagraphStyle(name="Body", parent=styles["Normal"], fontSize=11, leading=15, textColor=colors.HexColor("#0f172a")))
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=9, leading=12, textColor=colors.HexColor("#475569")))

    story: List[Any] = []
    story.append(Paragraph("IntelliDoc Answer Report", styles["AnswerHeading"]))
    story.append(Paragraph(f"Generated: {payload.generated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}", styles["MetaLabel"]))
    story.append(Spacer(1, 0.2 * inch))

    meta_rows = [
        [Paragraph("Document Name", styles["MetaLabel"]), Paragraph(_safe_text(payload.document_name), styles["Body"])],
        [Paragraph("Summary Type", styles["MetaLabel"]), Paragraph(payload.answer_mode, styles["Body"])],
        [Paragraph("Answer Length", styles["MetaLabel"]), Paragraph(payload.answer_length, styles["Body"])],
        [Paragraph("Pages", styles["MetaLabel"]), Paragraph(payload.page_range_label, styles["Body"])],
    ]
    meta_table = Table(meta_rows, colWidths=[1.4 * inch, 4.8 * inch])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.22 * inch))

    story.append(Paragraph("Question", styles["Heading2"]))
    story.append(Paragraph(payload.question.replace("\n", "<br/>"), styles["Body"]))
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph("AI Answer", styles["Heading2"]))
    answer_paragraphs = [p.strip() for p in payload.answer.split("\n") if p.strip()]
    for paragraph in answer_paragraphs:
        story.append(Paragraph(paragraph.replace("\n", "<br/>"), styles["Body"]))
        story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("Sources / Citations", styles["Heading2"]))

    if payload.sources:
        story.append(Paragraph("Sources", styles["Heading3"]))
        for source in payload.sources:
            story.append(Paragraph(f"• {_source_label(source)}", styles["Body"]))
        story.append(Spacer(1, 0.08 * inch))

    if payload.citations:
        story.append(Paragraph("Citations", styles["Heading3"]))
        for citation in payload.citations:
            story.append(Paragraph(f"• {_citation_label(citation)}", styles["Small"]))
        story.append(Spacer(1, 0.08 * inch))

    story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("Generated by IntelliDoc AI Document Intelligence Platform", styles["Small"]))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def generate_docx_bytes(payload: ExportPayload) -> bytes:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed")

    document = DocxDocument()

    title = document.add_paragraph()
    title_run = title.add_run("IntelliDoc Answer Report")
    title_run.bold = True
    title_run.font.size = Pt(18) if Pt else None

    meta = document.add_paragraph()
    meta.add_run(f"Generated: {payload.generated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}\n")
    meta.add_run(f"Document Name: {payload.document_name}\n")
    meta.add_run(f"Summary Type: {payload.answer_mode}\n")
    meta.add_run(f"Answer Length: {payload.answer_length}\n")
    meta.add_run(f"Pages: {payload.page_range_label}")

    document.add_heading("Question", level=2)
    document.add_paragraph(payload.question)

    document.add_heading("AI Answer", level=2)
    for paragraph in [p.strip() for p in payload.answer.split("\n") if p.strip()]:
        document.add_paragraph(paragraph)

    document.add_heading("Sources / Citations", level=2)
    if payload.sources:
        document.add_heading("Sources", level=3)
        for source in payload.sources:
            document.add_paragraph(_source_label(source), style="List Bullet")
    else:
        document.add_paragraph("No sources available.")

    if payload.citations:
        document.add_heading("Citations", level=3)
        for citation in payload.citations:
            document.add_paragraph(_citation_label(citation), style="List Bullet")

    document.add_paragraph("Generated by IntelliDoc AI Document Intelligence Platform")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_export_bytes(export_format: str, payload: ExportPayload) -> tuple[bytes, str]:
    normalized = (export_format or "").strip().lower()
    if normalized == "pdf":
        return generate_pdf_bytes(payload), "application/pdf"
    if normalized == "docx":
        return generate_docx_bytes(payload), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Unsupported export format: {export_format}")
